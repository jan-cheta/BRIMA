from base import Database
from model import Household, Resident, User, Blotter, Certificate, Barangay
from forms import (AddHouseholdForm, AddResidentForm, BrowseResidentForm,
    UpdateHouseholdForm, BrowseHouseholdForm, UpdateResidentForm, AddUserForm,
    UpdateUserForm, BrowseUserForm, AddBlotterForm, UpdateBlotterForm, BrowseBlotterForm,
    AddCertificateForm, UpdateCertificateForm, BrowseCertificateForm
)
from view import  BrimaView, MainView, LoginView
from widgets import BaseWindow, AboutWindow, SettingsWindow, DashboardWindow
from PySide6.QtWidgets import QMessageBox, QDialog, QFileDialog
from PySide6.QtCore import Qt, QDate, QSize
import pyqtgraph as pg
from pyqtgraph import PlotWidget, BarGraphItem, TextItem
from sqlalchemy import or_, and_, desc, select, create_engine, func
from sqlalchemy.orm import aliased, sessionmaker, declarative_base, Session             
from docx import Document
import os
import shutil
from datetime import datetime, date
import bcrypt
import pandas as pd



class MainController:
    def __init__(self, view: MainView):
        self.db = Database()
        self.session = self.db.get_session()
        self.view = view
        self.user = None
        self.brima = self.view.brima
        self.brima_control = BrimaController(self.brima, None)
        self.view.login.btLogin.clicked.connect(self.login)
        self.brima.btLogout.clicked.connect(self.logout)
        self.view.stack.setCurrentIndex(0)

    def login(self):
        """
        Login page logic for applying permissions and switching to BrimaView

        Returns:
            None
        """
        
        data = self.view.login.get_fields()
        user = self.is_valid_login(self.session, data)
        
        if user:
            QMessageBox.information(self.view, 'Success', 'Login Sucessful')
            self.view.login.tbUsername.clear()
            self.view.login.tbPassword.clear()
            self.user = user
            self.brima_control.user = self.user
            self.brima_control.set_titles()
            self.brima_control.apply_permissions()
            self.view.stack.setCurrentIndex(1)
            self.brima.stack.setCurrentIndex(7)
            self.view.showMaximized()
        else:
            QMessageBox.critical(self.view, 'Invalid Credentials', 'Please input valid credentials')

    def is_valid_login(self, session: Session, data: dict) -> User:
        """
        
        Login validation login that converts password to hash then checks the database for user

        Args:
            session (Session): The session to be used to query the database for users
            data (dict): The data from the fields of the login page

        Returns:
            User: returns the user if query is valid 
            
        """        
        user = session.query(User).filter(User.username==data.get('username')).first()

        if user and bcrypt.checkpw(data.get('password').encode('utf-8'), user.password):
            return user
        else:
            return None

    def logout(self):
        self.brima_control.user = None
        self.view.stack.setCurrentIndex(0)
        self.view.showNormal()
    
    
class BrimaController:
    def __init__(self, view : BrimaView, user):
        self.db = Database()
        self.session = self.db.get_session()
        self.view = view
        self.user = user 

        self.household_control = HouseholdWindowController(self.view.household_window)
        self.resident_control = ResidentWindowController(self.view.resident_window)
        self.user_control = UserWindowController(self.view.admin_window)
        self.blotter_control = BlotterWindowController(self.view.blotter_window)
        self.certificate_control = CertificateWindowController(self.view.certificate_window)
        self.about_control = AboutUsWindowController(self.view.about_window)
        self.settings_control = SettingsWindowController(self.view.settings_window)
        self.dashboard_control = DashboardWindowController(self.view.dashboard_window)

        self.view.btHousehold.clicked.connect(lambda: self.view.stack.setCurrentIndex(0))
        self.view.btResident.clicked.connect(lambda: self.view.stack.setCurrentIndex(1))
        self.view.btAdmin.clicked.connect(lambda: self.view.stack.setCurrentIndex(2))
        self.view.btBlotter.clicked.connect(lambda: self.view.stack.setCurrentIndex(3))
        self.view.btCertificate.clicked.connect(lambda: self.view.stack.setCurrentIndex(4))
        self.view.btAboutUs.clicked.connect(lambda: self.view.stack.setCurrentIndex(5))
        self.view.btAboutUs.clicked.connect(self.about_control.load_data)
        self.view.btSettings.clicked.connect(lambda: self.view.stack.setCurrentIndex(6))
        self.view.btDashboard.clicked.connect(lambda: self.view.stack.setCurrentIndex(7))
        self.view.btDashboard.clicked.connect(self.dashboard_control.load_data)
        self.view.stack.setCurrentIndex(7)

    
    def set_titles(self):
        self.view.lbUser.setText(f"HELLO, {self.user.resident.first_name} {self.user.resident.last_name}!")
        barangay = self.session.query(Barangay).first()
        self.view.lbBrgy.setText(f"{barangay.name}")
    
    def apply_permissions(self):
        is_admin = (self.user.position in ['CAPTAIN', 'SECRETARY'])

        self.view.btAdmin.setVisible(is_admin)

        if not is_admin:
            self.setup_admin(self.view)
        else:
            self.setup_user(self.view)

    def setup_admin(self, view: BrimaView) -> None:
        """
        
        Sets up the all widgets to visible 

        Args:
            view (BrimaView): The main app UI

        Returns:
            None
        """
        for win in (
            view.household_window,
            view.resident_window,
            view.blotter_window,
            view.certificate_window,
            view.admin_window,
        ):
            for btn in (win.btAdd, win.btEdit, win.btDelete):
                btn.setVisible(False)

        for btn in (view.settings_window.edit_barangay, view.settings_window.backup):
            btn.setEnabled(False)

    def setup_user(self, view: BrimaView) -> None:
        """
        
        Sets admin level widgets to not visible

        Args:
            view (BrimaView): The main app UI

        Returns:
            None
        """
        for win in (
            view.household_window,
            view.resident_window,
            view.blotter_window,
            view.certificate_window,
            view.admin_window,
        ):
            for btn in (win.btAdd, win.btEdit, win.btDelete):
                btn.setVisible(True)
    
        for btn in (view.settings_window.edit_barangay, view.settings_window.backup):
            btn.setEnabled(True)

class HouseholdWindowController:
    def __init__(self, view: BaseWindow):
        self.db = Database()
        self.session = self.db.get_session()
        self.view = view
        self.refresh()
        
        self.view.btRefresh.clicked.connect(self.refresh)
        self.view.btSearch.clicked.connect(self.search)
        self.view.tbSearchBar.returnPressed.connect(self.search)
        self.view.btAdd.clicked.connect(self.add)
        self.view.btEdit.clicked.connect(self.edit)
        self.view.btDelete.clicked.connect(self.delete)
        self.view.btBrowse.clicked.connect(self.browse)
        
    def refresh(self):
        self.view.set_search_text('')
        households = self.session.query(Household).order_by(Household.household_name).all()
        data = []
        for household in households:
            result = [
                household.id,
                household.date_added,
                household.household_name,
                f"{household.house_no} {household.street} {household.sitio} {household.landmark}"
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Date Added', 'Household Name', 'Address'], data)
    
    def search(self):
        search_text = self.view.get_search_text().lower()  
        search_terms = search_text.split() 
        
        # Base query
        query = self.session.query(Household)
        
        # Apply AND of ORs: each term must match one of the fields
        if search_terms:
            conditions = []
            for term in search_terms:
                term_filter = or_(
                    Household.date_added.ilike(f"%{term}%"),
                    Household.household_name.ilike(f"%{term}%"),
                    Household.house_no.ilike(f"%{term}%"),
                    Household.street.ilike(f"%{term}%"),
                    Household.sitio.ilike(f"%{term}%"),
                    Household.landmark.ilike(f"%{term}%")
                )
                conditions.append(term_filter)
            
            query = query.filter(and_(*conditions))
        
        # Order and execute
        households = query.order_by(Household.household_name).all()
        
        data = []
        for household in households:
            address = " ".join(filter(None, [
                household.house_no,
                household.street,
                household.sitio,
                household.landmark
            ]))
            
            result = [
                household.id,
                household.date_added,
                household.household_name,
                address
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Date Added', 'Household Name', 'Address'], data)
    
    def add(self):
            add_form = AddHouseholdForm()
        
            # Button signals
            add_form.addbar.btAdd.clicked.connnect(lambda: self.on_add_button_click(add_form))
            add_form.addbar.btCancel.clicked.connect(add_form.reject)
        
            # Execute the form as a modal dialog
            add_form.exec()
    
    def on_add_button_click(self, add_form):
        data = add_form.get_fields()

        validate = self.validate_addlidate_add(self.session, data)
        
        if not validate[0]:
            QMessageBox.critical(add_form, "Error", validate[1])
        
        try:
            # Add the new resident to the session and commit the transaction
            self.session.add(self.newhousehold_make(data))
            self.session.commit()
            QMessageBox.information(add_form, "Success", "Household added successfully!")
            self.refresh() 
            add_form.accept()
        except Exception as e:
            self.session.rollback()
            QMessageBox.critical(add_form, "Error", f"Failed to add household: {str(e)}")

    def validate_add(session: Session, data: dict) -> tuple[bool, str]:
        """
    
        Validation logic for adding household

        Args:
            session (Session): The session to be used to query the database.
            data (dict): The data from the fields.

        Returns:
            tuple[bool,str]: boolean to show validity and error message if not valid
        """ 
        household_name = data.get("household_name")

        # Check if household_name is null
        if not household_name:
             return False, "Household name cannot be empty."
    
        # Check if household already exists
        exist_household = session.query(Household).filter(func.upper(Household.household_name) == household_name.upper()).first()
        if exist_household:
             return False, "Household name already exists."
        return True, ""

    def newhousehold_make(data: dict) -> Household:
        """
    
        Creates a new Household entity from a data dictionary

        Args:
            data (dict): Household data to be used for new entity

        Returns:
            Household: New Household entity
        """
        return Household(
            household_name = data.get("household_name", "").upper(),
            house_no = data.get("house_no", "").upper(),
            street = data.get("street", "").upper(),
            sitio = data.get("sitio", "").upper(),
            landmark = data.get("landmark", "").upper()
         )

    
    def edit(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Household to Edit')
            return
        
        if row_id:
            household = self.session.query(Household).get(row_id)
            
            if household:
                update_form = UpdateHouseholdForm()
                
                update_form.set_fields(
                    household_name=household.household_name,
                    house_no=household.house_no,
                    street=household.street,
                    sitio=household.sitio,
                    landmark=household.landmark
                )
                
                update_form.updatebar.btRevert.clicked.connect(
                    lambda: update_form.set_fields(
                        household_name=household.household_name,
                        house_no=household.house_no,
                        street=household.street,
                        sitio=household.sitio,
                        landmark=household.landmark
                    )
                )
                
                update_form.updatebar.btUpdate.clicked.connect(
                    lambda: self.save_update(household, update_form)
                )
                
                update_form.updatebar.btCancel.clicked.connect(update_form.reject)
                
                if update_form.exec() == QDialog.Accepted:
                    self.refresh()
    
    def save_update(self, household, update_form):
        data = update_form.get_fields()
        data = {key: value.upper() if isinstance(value, str) else value for key, value in data.items()}
        
        exist_household = self.session.query(Household).filter(func.upper(Household.household_name) == data.get("household_name").upper()).all()
        
        if not data.get("household_name"):  
            QMessageBox.critical(update_form, "Error", "Household name cannot be empty.")
            return

        if not household.household_name == data.get("household_name"):
            if exist_household:
                QMessageBox.critical(update_form, "Error", "Household name already exists.")
                return

        household.household_name = data['household_name'].upper()
        household.house_no = data['house_no'].upper()
        household.street = data['street'].upper()
        household.sitio = data['sitio'].upper()
        household.landmark = data['landmark'].upper()
        
        try:
            self.session.commit()
            QMessageBox.information(self.view, "Success", "Household updated successfully!")
        except Exception as e:
            self.session.rollback()  
            QMessageBox.critical(self.view, "Error", f"Failed to update household: {str(e)}")
        
        update_form.accept()   
        
    def browse(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Household to Browse')
            return
        if row_id:
            household = self.session.query(Household).get(row_id)
            
            if household:
                browse_form = BrowseHouseholdForm()
                residents = household.residents

                browse_form.set_fields(
                    household_name=household.household_name,
                    house_no=household.house_no,
                    street=household.street,
                    sitio=household.sitio,
                    landmark=household.landmark
                )
                
                data = []
                for resident in residents:
                    full_name = " ".join(filter(None, [resident.first_name, resident.middle_name, resident.suffix]))
                    full_name = f"{resident.last_name}, {full_name}".strip()

                    result = [
                        resident.id,
                        resident.date_added,
                        full_name,
                        resident.role,
                    ]
                    data.append(result)
                
                browse_form.load_table(['id', 'Date Added', 'Household Member', 'Role'], data)
                browse_form.exec()
        
    
    def delete(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Household to Delete')
            return
        if row_id:
            household = self.session.query(Household).get(row_id)
            if household:
                reply = QMessageBox.question(
                    self.view,  # Parent widget
                    "Confirm Deletion",
                    f"Are you sure you want to delete '{household.household_name}'?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
    
                if reply == QMessageBox.Yes:
                    self.session.delete(household)
                    self.session.commit()
                    self.refresh()

class ResidentWindowController:
    def __init__(self, view: BaseWindow):
        self.db = Database()
        self.session = self.db.get_session()
        self.view = view
        self.refresh()
        
        self.view.btRefresh.clicked.connect(self.refresh)
        self.view.btSearch.clicked.connect(self.search)
        self.view.tbSearchBar.returnPressed.connect(self.search)
        self.view.btAdd.clicked.connect(self.add)
        self.view.btEdit.clicked.connect(self.edit)
        self.view.btDelete.clicked.connect(self.delete)
        self.view.btBrowse.clicked.connect(self.browse)
        
    def refresh(self):
        self.view.set_search_text('')
        residents = self.session.query(Resident).order_by(Resident.last_name).all()
        data = []
        for resident in residents:
            full_name = " ".join(filter(None, [resident.first_name, resident.middle_name, resident.suffix]))
            full_name = f"{resident.last_name}, {full_name}".strip()

            household = resident.household
            household_name = household.household_name if household else "N/A"
            address_parts = [
                getattr(household, "house_no", ""),
                getattr(household, "street", ""),
                getattr(household, "sitio", ""),
                getattr(household, "landmark", "")
            ] if household else []

            address = " ".join(filter(None, address_parts))

            result = [
                resident.id,
                resident.date_added,
                full_name,
                resident.role,
                household_name,
                address
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Date Added', 'Full Name', 'Role', 'Household Name', 'Address'], data)

    def search(self):
        search_text = self.view.get_search_text().lower()
        search_terms = search_text.split()

        query = self.session.query(Resident).join(Resident.household)

        if search_terms:
            conditions = []
            for term in search_terms:
                term_filter = or_(
                    Resident.date_added.ilike(f"%{term}%"),
                    Resident.first_name.ilike(f"%{term}%"),
                    Resident.middle_name.ilike(f"%{term}%"),
                    Resident.last_name.ilike(f"%{term}%"),
                    Resident.suffix.ilike(f"%{term}%"),
                    Resident.role.ilike(f"%{term}%"),
                    Household.household_name.ilike(f"%{term}%"),
                    Household.house_no.ilike(f"%{term}%"),
                    Household.street.ilike(f"%{term}%"),
                    Household.sitio.ilike(f"%{term}%"),
                    Household.landmark.ilike(f"%{term}%")
                )
                conditions.append(term_filter)
            
            query = query.filter(and_(*conditions))

        residents = query.order_by(Resident.last_name).all()
        data = []
        for resident in residents:
            full_name = " ".join(filter(None, [
                resident.first_name,
                resident.middle_name,
                resident.suffix
            ]))
            full_name = f"{resident.last_name}, {full_name}".strip()

            household = resident.household
            household_name = household.household_name if household else "N/A"
            address_parts = [
                getattr(household, "house_no", ""),
                getattr(household, "street", ""),
                getattr(household, "sitio", ""),
                getattr(household, "landmark", "")
            ] if household else []

            address = " ".join(filter(None, address_parts))

            result = [
                resident.id,
                resident.date_added,
                full_name,
                resident.role,
                household_name,
                address
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Date Added', 'Full Name', 'Role', 'Household Name', 'Address'], data)
    
    def add(self):
        add_form = AddResidentForm()
    
        # Populate households for dropdown (household names and IDs)
        households = self.session.query(Household).order_by(Household.household_name).all()
        household_dict = {household.household_name: household.id for household in households}
    
        # Set the fields for the form (populate the household dropdown with names)
        add_form.set_fields(household=list(household_dict.keys()))
        add_form.form.cbHousehold.setCurrentText('')  # Ensure no default value is set
        add_form.form.cbHousehold.currentTextChanged.connect(lambda: self.autofill_household(household_dict, add_form))
    
        # Button signals
        add_form.addbar.btAdd.clicked.connect(lambda: self.on_add_button_click(add_form, household_dict))
        add_form.addbar.btCancel.clicked.connect(add_form.reject)
    
        # Execute the form as a modal dialog
        add_form.exec()
    
    def on_add_button_click(self, add_form, household_dict):
        # Get the data from the form
        data = add_form.get_fields()
    
        # Check if 'household' is present and valid in the data
        household_name = data.get("household")  # Use get to avoid KeyError
        if not household_name:  # Check if the household field is empty
            QMessageBox.critical(add_form, "Error", "Please select a Household.")
            return  # Exit if the household is not selected or invalid
        
        # Retrieve the corresponding Household ID from the dictionary
        household_id = household_dict.get(household_name)
        if not household_id:
            QMessageBox.critical(add_form, "Error", f"No matching Household found for '{household_name}'.")
            return  # Exit if no valid household ID is found
        
        conditions = and_(
            func.upper(Resident.first_name) == data.get("first_name").upper(),
            func.upper(Resident.last_name) == data.get("last_name").upper(),
            func.upper(Resident.middle_name) == data.get("middle_name").upper(),
            func.upper(Resident.suffix) == data.get("suffix").upper()
        )

        exist_resident = self.session.query(Resident).filter(conditions).first()
        if exist_resident:
            QMessageBox.critical(add_form, "Error", "Resident already exists")
            return
        
        if not data.get("first_name") or not data.get("last_name"):
            QMessageBox.critical(add_form, "Error", "First name and Last name cannot be empty")
            return
        
        if not data.get("citizenship"):
            QMessageBox.critical(add_form, "Error", "Citizenship cannot be empty")
            return

        # Get the Household object from the database using the household ID
        household = self.session.query(Household).get(household_id)
    
        # Remove the 'household' key from the data dictionary since it's already handled
        del data['household']
    
        # Uppercase all string values in data (excluding non-string types)
        data = {key: value.upper() if isinstance(value, str) else value for key, value in data.items()}
    
        # Create a new Resident instance with the provided data
        new_resident = Resident(**data)
        new_resident.household = household  # Associate the resident with the selected household
    
        try:
            # Add the new resident to the session and commit the transaction
            self.session.add(new_resident)
            self.session.commit()
            QMessageBox.information(add_form, "Success", "Resident added successfully!")
            self.refresh()  # Refresh the view/list of residents
            
            # Close the dialog only if everything is successful
            add_form.accept()
        except Exception as e:
            # If any error occurs, roll back the session and display an error message
            self.session.rollback()
            QMessageBox.critical(add_form, "Error", f"Failed to add resident: {str(e)}")
            add_form.activateWindow()  # Ensure focus returns to the dialog
            add_form.raise_()
    
    def autofill_household(self, data, view):
        id = data[view.form.cbHousehold.currentText()]
        
        if id:
            household = self.session.query(Household).get(id)
            
            view.form.tbHouseholdName.setText(household.household_name)
            view.form.tbHouseNo.setText(household.house_no)
            view.form.tbStreet.setText(household.street)
            view.form.cbSitio.setCurrentText(household.sitio)
            view.form.tbLandmark.setText(household.landmark)
        
    
    def edit(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Resident to Edit')
            return
        if row_id:
            resident = self.session.query(Resident).get(row_id)
    
            if resident:
                update_form = UpdateResidentForm()
    
                # Fetch households
                households = self.session.query(Household).order_by(Household.household_name).all()
                household_dict = {household.household_name: household.id for household in households}
                household_names = list(household_dict.keys())
    
                # Populate household combo box
                update_form.form.cbHousehold.clear()
                update_form.form.cbHousehold.addItems(household_names)
                update_form.form.cbHousehold.setCurrentText(resident.household.household_name if resident.household else '')
                self.autofill_household(household_dict, update_form)
    
                update_form.form.cbHousehold.currentTextChanged.connect(
                    lambda: self.autofill_household(household_dict, update_form)
                )
    
                # Set initial values
                update_form.set_fields(
                    first_name=resident.first_name,
                    last_name=resident.last_name,
                    middle_name=resident.middle_name,
                    suffix=resident.suffix,
                    date_of_birth=resident.date_of_birth,
                    phone1=resident.phone1,
                    phone2=resident.phone2,
                    email=resident.email,
                    household=resident.household.household_name if resident.household else '',
                    occupation=resident.occupation,
                    civil_status=resident.civil_status,
                    citizenship=resident.citizenship,
                    education=resident.education,
                    remarks=resident.remarks,
                    sex=resident.sex,
                    role=resident.role,
                )
    
                # Revert button restores original values
                update_form.updatebar.btRevert.clicked.connect(
                    lambda: update_form.set_fields(
                        first_name=resident.first_name,
                        last_name=resident.last_name,
                        middle_name=resident.middle_name,
                        suffix=resident.suffix,
                        date_of_birth=resident.date_of_birth,
                        phone1=resident.phone1,
                        phone2=resident.phone2,
                        email=resident.email,
                        household=resident.household.household_name if resident.household else '',
                        occupation=resident.occupation,
                        civil_status=resident.civil_status,
                        citizenship=resident.citizenship,
                        education=resident.education,
                        remarks=resident.remarks,
                        sex=resident.sex,
                        role=resident.role,
                    )
                )
    
                # Save button
                update_form.updatebar.btUpdate.clicked.connect(
                    lambda: self.save_update(resident, update_form, household_dict)
                )
    
                # Cancel button
                update_form.updatebar.btCancel.clicked.connect(update_form.reject)
    
                update_form.exec()
    
    
    def save_update(self, resident, update_form, household_dict):
        updated_data = update_form.get_fields()
    
        # Map household name to ID
        household_name = update_form.form.cbHousehold.currentText()
        household_id = household_dict.get(household_name)
    
        if not household_id:
            QMessageBox.critical(update_form, "Error", "Please select a valid household.")
            return
    
        household = self.session.query(Household).get(household_id)
        if not household:
            QMessageBox.critical(update_form, "Error", "Selected household does not exist.")
            return
        
        conditions = and_(
            func.upper(Resident.first_name) == updated_data.get("first_name").upper(),
            func.upper(Resident.last_name) == updated_data.get("last_name").upper(),
            func.upper(Resident.middle_name) == updated_data.get("middle_name").upper(),
            func.upper(Resident.suffix) == updated_data.get("suffix").upper()
        )

        exist_resident = self.session.query(Resident).filter(conditions).all()

        con = [
            resident.first_name.upper() == updated_data.get("first_name").upper(),
            resident.last_name.upper() == updated_data.get("last_name").upper(),
            resident.middle_name.upper() == updated_data.get("middle_name").upper(),
            resident.suffix.upper() == updated_data.get("suffix").upper()
        ]
        
        if not all(con):
            if exist_resident:
                QMessageBox.critical(update_form, "Error", "Resident already exists")
                return
        
        if not updated_data.get("first_name") or not updated_data.get("last_name"):
            QMessageBox.critical(update_form, "Error", "First name and Last name cannot be empty")
            return
        
        if not updated_data.get("citizenship"):
            QMessageBox.critical(update_form, "Error", "Citizenship cannot be empty")
            return
    
        # Update resident fields
        resident.first_name = updated_data["first_name"].upper()
        resident.last_name = updated_data["last_name"].upper()
        resident.middle_name = updated_data["middle_name"].upper()
        resident.suffix = updated_data["suffix"].upper()
        resident.date_of_birth = updated_data["date_of_birth"]
        resident.phone1 = updated_data["phone1"]
        resident.phone2 = updated_data["phone2"]
        resident.email = updated_data["email"]
        resident.occupation = updated_data["occupation"].upper()
        resident.civil_status = updated_data["civil_status"].upper()
        resident.citizenship = updated_data['citizenship'].upper()
        resident.education = updated_data["education"].upper()
        resident.remarks = updated_data["remarks"].upper()
        resident.sex = updated_data["sex"].upper()
        resident.role = updated_data["role"].upper()
        resident.household = household
    
        try:
            self.session.commit()
            QMessageBox.information(update_form, "Success", "Resident updated successfully!")
            update_form.accept()
            self.refresh()
        except Exception as e:
            self.session.rollback()
            QMessageBox.critical(update_form, "Error", f"Failed to update resident: {str(e)}")
        
    def browse(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Resident to Browse')
            return
        
        if row_id:
            resident = self.session.query(Resident).get(row_id)
            
            if resident:
                browse_form = BrowseResidentForm()
                
                browse_form.set_fields(
                    first_name=resident.first_name,
                    last_name=resident.last_name,
                    middle_name=resident.middle_name,
                    suffix=resident.suffix,
                    date_of_birth=resident.date_of_birth,
                    phone1=resident.phone1,
                    phone2=resident.phone2,
                    email=resident.email,
                    household=resident.household.household_name if resident.household else '',
                    occupation=resident.occupation,
                    civil_status=resident.civil_status,
                    citizenship=resident.citizenship,
                    education=resident.education,
                    remarks=resident.remarks,
                    sex=resident.sex,
                    role=resident.role,
                )
                
                households = self.session.query(Household).order_by(Household.household_name).all()
                household_dict = {household.household_name: household.id for household in households}
                household_names = list(household_dict.keys())
    
                # Populate household combo box
                browse_form.form.cbHousehold.setCurrentText(resident.household.household_name if resident.household else '')
                self.autofill_household(household_dict, browse_form)
                
                browse_form.exec()
        
    
    def delete(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Resident to Delete')
            return
        if row_id:
            resident = self.session.query(Resident).get(row_id)
            if resident:
                reply = QMessageBox.question(
                    self.view, 
                    "Confirm Deletion",
                    f"Are you sure you want to delete {resident.last_name}, {resident.first_name} {resident.middle_name} {resident.suffix}?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
    
                if reply == QMessageBox.Yes:
                    self.session.delete(resident)
                    self.session.commit()
                    self.refresh()

class UserWindowController:
    def __init__(self, view: BaseWindow):
        self.db = Database()
        self.session = self.db.get_session()
        self.view = view
        self.refresh()
        
        self.view.btRefresh.clicked.connect(self.refresh)
        self.view.btSearch.clicked.connect(self.search)
        self.view.tbSearchBar.returnPressed.connect(self.search)
        self.view.btAdd.clicked.connect(self.add)
        self.view.btEdit.clicked.connect(self.edit)
        self.view.btDelete.clicked.connect(self.delete)
        self.view.btBrowse.clicked.connect(self.browse)
    
    def refresh(self):
        self.view.set_search_text('')
        users = self.session.query(User).join(User.resident).order_by(User.id, User.date_added).all()
        data = []
        for user in users:
            username = user.username
            position = user.position
            resident = user.resident
            resident_name = [
                getattr(resident, "first_name", ""),
                getattr(resident, "middle_name", ""),
                getattr(resident, "last_name", ""),
                getattr(resident, "suffix", "")
            ] if resident else []

            full_name = " ".join(filter(None, resident_name))

            result = [
                user.id,
                user.date_added,
                username,
                position,
                full_name
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Date Added', 'Username', 'Position', 'Full Name'], data)

    def search(self):
        search_text = self.view.get_search_text().lower()
        search_terms = search_text.split()

        query = self.session.query(User).join(User.resident)

        if search_terms:
            conditions = []
            for term in search_terms:
                term_filter = or_(
                    User.date_added.ilike(f"%{term}%"),
                    User.username.ilike(f"%{term}%"),
                    User.position.ilike(f"%{term}%"),
                    Resident.first_name.ilike(f"%{term}%"),
                    Resident.middle_name.ilike(f"%{term}%"),
                    Resident.last_name.ilike(f"%{term}%"),
                    Resident.suffix.ilike(f"%{term}%")
                )
                conditions.append(term_filter)
            
            query = query.filter(and_(*conditions))

        users = query.order_by(User.id, User.date_added).all()
        data = []
        for user in users:
            username = user.username
            position = user.position
            resident = user.resident
            resident_name = [
                getattr(resident, "first_name", ""),
                getattr(resident, "middle_name", ""),
                getattr(resident, "last_name", ""),
                getattr(resident, "suffix", "")
            ] if resident else []

            full_name = " ".join(filter(None, resident_name))

            result = [
                user.id,
                user.date_added,
                username,
                position,
                full_name
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Date Added', 'Username', 'Position', 'Full Name'], data)
    
    def add(self):
        add_form = AddUserForm()
    
        # Populate households for dropdown (household names and IDs)
        residents = self.session.query(Resident).filter(Resident.user == None).order_by(Resident.last_name).all()
        resident_dict = {
            " ".join(filter(None, [
                resident.first_name,
                resident.middle_name,
                resident.last_name,
                resident.suffix
            ])): resident.id
            for resident in residents
        }
    
        # Set the fields for the form (populate the household dropdown with names)
        add_form.set_fields(name=list(resident_dict.keys()))
        add_form.form.cbName.setCurrentText('')  # Ensure no default value is set
    
        # Button signals
        add_form.addbar.btAdd.clicked.connect(lambda: self.on_add_button_click(add_form, resident_dict))
        add_form.addbar.btCancel.clicked.connect(add_form.reject)
    
        # Execute the form as a modal dialog
        add_form.exec()
    
    def on_add_button_click(self, add_form, resident_dict):
        data = add_form.get_fields()

        if not data.get('username'):
            QMessageBox.critical(add_form, 'Error', 'Username cannot be empty')
            return
        
        exist_username = self.session.query(User).filter(User.username == data.get('username')).first()

        if exist_username:
            QMessageBox.critical(add_form, 'Error', 'Username already taken')
            return
        
        if not data.get('password'):
            QMessageBox.critical(add_form, 'Error', 'Password cannot be empty')
            return
        
        if not data.get('confirm_password') == data.get('password'):
            QMessageBox.critical(add_form, 'Error', 'Password do not match')
    
        resident_name = data.get("name")
        if not resident_name:  
            QMessageBox.critical(add_form, "Error", "Please select a Resident.")
            return  
        
        resident_id = resident_dict.get(resident_name)
        if not resident_id:
            QMessageBox.critical(add_form, "Error", f"No matching Household found for '{resident_name}'.")
            return  
            
        resident = self.session.query(Resident).get(resident_id)
    
        # Remove the 'household' key from the data dictionary since it's already handled
        del data['name']

        del data['confirm_password']

        salt = bcrypt.gensalt()
        data['password'] = bcrypt.hashpw(data.get('password').encode('utf-8'), salt)
        
        # Uppercase all string values in data (excluding non-string types)
        data = {key: value.upper() if isinstance(value, str) and key != 'username' and key != 'password' else value for key, value in data.items()}
    
        # Create a new Resident instance with the provided data
        new_user = User(**data)
        new_user.resident = resident  # Associate the resident with the selected household
    
        try:
            # Add the new resident to the session and commit the transaction
            self.session.add(new_user)
            self.session.commit()
            QMessageBox.information(add_form, "Success", "User added successfully!")
            self.refresh()  
            
            # Close the dialog only if everything is successful
            add_form.accept()
        except Exception as e:
            self.session.rollback()
            QMessageBox.critical(add_form, "Error", f"Failed to add resident: {str(e)}")
            add_form.activateWindow() 
            add_form.raise_()    
    
    def edit(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select User to Edit')
            return
        if row_id:
            user = self.session.query(User).get(row_id)
    
            if user:
                update_form = UpdateUserForm()
    
                # Fetch households
                residents = self.session.query(Resident).order_by(Resident.last_name).all()
                resident_dict = {
                    " ".join(filter(None, [
                        resident.first_name,
                        resident.middle_name,
                        resident.last_name,
                        resident.suffix
                    ])): resident.id
                    for resident in residents
                }
                resident_names = list(resident_dict.keys())
    
                # Populate household combo box
                update_form.form.cbName.clear()
                update_form.form.cbName.addItems(resident_names)
                full_name = ""
                if user.resident:
                    full_name = " ".join(filter(None, [
                        user.resident.first_name,
                        user.resident.middle_name,
                        user.resident.last_name,
                        user.resident.suffix
                    ]))
                    update_form.form.cbName.setCurrentText(full_name)
    
    
                # Set initial values
                update_form.set_fields(
                    name = full_name,
                    username = user.username,
                    position = user.position
                )
    
                # Revert button restores original values
                update_form.updatebar.btRevert.clicked.connect(
                    lambda: update_form.set_fields(
                        name = full_name,
                        username = user.username,
                        position = user.position
                    )
                )
    
                # Save button
                update_form.updatebar.btUpdate.clicked.connect(
                    lambda: self.save_update(user, update_form, resident_dict)
                )
    
                # Cancel button
                update_form.updatebar.btCancel.clicked.connect(update_form.reject)
    
                update_form.exec()
    
    
    def save_update(self, user, update_form, resident_dict):
        updated_data = update_form.get_fields()
    
        # Map household name to ID
        resident_name = update_form.form.cbName.currentText()
        resident_id = resident_dict.get(resident_name)


        if not updated_data.get('username'):
            QMessageBox.critical(update_form, 'Error', 'Username cannot be empty')
            return
        
        if not updated_data.get('username') == user.username:
            exist_username = self.session.query(User).filter(User.username == updated_data.get('username')).first()

            if exist_username:
                QMessageBox.critical(update_form, 'Error', 'Username already taken')
                return
        
        if not updated_data.get('password'):
            QMessageBox.critical(update_form, 'Error', 'Password cannot be empty')
            return
        
        if not updated_data.get('confirm_password') == updated_data.get('password'):
            QMessageBox.critical(update_form, 'Error', 'Password do not match')
            return

        if not resident_name:
            QMessageBox.critical(update_form, "Error", "Please select a valid resident.")
            return
    
        resident = self.session.query(Resident).get(resident_id)
        if not resident:
            QMessageBox.critical(update_form, "Error", "Selected resident does not exist.")
            return
    
        # Update resident fields
        user.username = updated_data['username']
        user.password = bcrypt.hash(updated_data['password'])
        user.position = updated_data['position']
        user.resident = resident
    
        try:
            self.session.commit()
            QMessageBox.information(update_form, "Success", "User updated successfully!")
            update_form.accept()
            self.refresh()
        except Exception as e:
            self.session.rollback()
            QMessageBox.critical(update_form, "Error", f"Failed to update User: {str(e)}")
        
    def browse(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select User to Browse')
            return
        if row_id:
            user = self.session.query(User).get(row_id)
            
            if user:
                browse_form = BrowseUserForm()
                resident = user.resident
                
                full_name = " ".join(filter(None, [
                    resident.first_name,
                    resident.middle_name,
                    resident.last_name,
                    resident.suffix
                ]))
                browse_form.set_fields(
                    name = full_name,
                    username = user.username,
                    position = user.position
                   
                )
                
                browse_form.exec()
    
    def delete(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select User to Delete')
            return
        if row_id:
            user = self.session.query(User).get(row_id)
            if user:
                reply = QMessageBox.question(
                    self.view, 
                    "Confirm Deletion",
                    f"Are you sure you want to delete User: {user.username}?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
    
                if reply == QMessageBox.Yes:
                    self.session.delete(user)
                    self.session.commit()
                    self.refresh()

class BlotterWindowController:
    def __init__(self, view : BaseWindow):
        self.db = Database()
        self.session = self.db.get_session()
        self.view = view
        self.refresh()
        
        self.view.btRefresh.clicked.connect(self.refresh)
        self.view.btSearch.clicked.connect(self.search)
        self.view.tbSearchBar.returnPressed.connect(self.search)
        self.view.btAdd.clicked.connect(self.add)
        self.view.btEdit.clicked.connect(self.edit)
        self.view.btDelete.clicked.connect(self.delete)
        self.view.btBrowse.clicked.connect(self.browse)
    
    def refresh(self):
        self.view.set_search_text('')
        blotters = self.session.query(Blotter).order_by(desc(Blotter.record_date)).all()
        data = []
        for blotter in blotters:
            result = [
                blotter.id,
                blotter.record_date,
                blotter.complainant,
                blotter.respondent,
                blotter.status,
                str(blotter.full_report)[:20]
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Record Date', 'Complainant', 'Respondent', 'Status','Report'], data)
    
    def search(self):
        search_text = self.view.get_search_text().lower()  
        search_terms = search_text.split() 
        
        # Base query
        query = self.session.query(Blotter)
        
        # Apply AND of ORs: each term must match one of the fields
        if search_terms:
            conditions = []
            for term in search_terms:
                term_filter = or_(
                    Blotter.record_date.ilike(f"%{term}%"),
                    Blotter.complainant.ilike(f"%{term}%"),
                    Blotter.respondent.ilike(f"%{term}%"),
                    Blotter.status.ilike(f"%{term}%"),
                    Blotter.full_report.ilike(f"%{term}%"),
                )
                conditions.append(term_filter)
            
            query = query.filter(and_(*conditions))
        
        # Order and execute
        blotters = query.order_by(desc(Blotter.record_date)).all()
        
        data = []
        for blotter in blotters:
            result = [
                blotter.id,
                blotter.record_date,
                blotter.complainant,
                blotter.respondent,
                blotter.status,
                str(blotter.full_report)[:20]
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Record Date', 'Complainant', 'Respondent', 'Status','Report'], data)
    
    def add(self):
            add_form = AddBlotterForm()
            add_form.addbar.btAdd.clicked.connect(lambda: self.on_add_button_click(add_form))
            add_form.addbar.btCancel.clicked.connect(add_form.reject)
            add_form.form.tbRecordDate.setDate(QDate.currentDate())
            
            # Execute the form as a modal dialog
            add_form.exec()

    def on_add_button_click(self, add_form):
        data = add_form.get_fields()
        data = {key: value.upper() if isinstance(value, str) else value for key, value in data.items()}
        
        if not data.get('nature_of_dispute'):
            QMessageBox.critical(add_form ,'Error', 'Nature of Dispute cannot be empty')
            return

        if not data.get('complainant'):
            QMessageBox.critical(add_form ,'Error', 'Complainant cannot be empty')
            return
        
        if not data.get('respondent'):
            QMessageBox.critical(add_form ,'Error', 'Respondent cannot be empty')
            return
        
        if not data.get('full_report'):
            QMessageBox.critical(add_form ,'Error', 'Report cannot be empty')
            return

        new_blotter = Blotter(**data)
        
        try:
            self.session.add(new_blotter)
            self.session.commit()
            QMessageBox.information(add_form, "Success", "Blotter added successfully!")
            add_form.accept()
            self.refresh()
        except Exception as e:
            self.session.rollback()
            QMessageBox.critical(add_form, "Error", f"Failed to add blotter: {str(e)}")
        
    def edit(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Blotter to Edit')
            return
        if row_id:
            blotter = self.session.query(Blotter).get(row_id)
            
            if blotter:
                update_form = UpdateBlotterForm()
                
                update_form.set_fields(
                    record_date = blotter.record_date,
                    status = blotter.status,
                    action_taken = blotter.action_taken,
                    nature_of_dispute = blotter.nature_of_dispute,
                    complainant = blotter.complainant,
                    respondent = blotter.respondent,
                    full_report = blotter.full_report
                )
                
                update_form.updatebar.btRevert.clicked.connect(
                    lambda: update_form.set_fields(
                        record_date = blotter.record_date,
                        status = blotter.statis,
                        action_taken = blotter.action_taken,
                        nature_of_dispute = blotter.nature_of_dispute,
                        complainant = blotter.complainant,
                        respondent = blotter.respondent,
                        full_report = blotter.full_report
                    )
                )
                
                update_form.updatebar.btUpdate.clicked.connect(
                    lambda: self.save_update(blotter, update_form)
                )
                
                update_form.updatebar.btCancel.clicked.connect(update_form.reject)
                
                if update_form.exec() == QDialog.Accepted:
                    self.refresh()
    
    def save_update(self, blotter, update_form):
        data = update_form.get_fields()
        data = {key: value.upper() if isinstance(value, str) else value for key, value in data.items()}
        
        if not data.get('nature_of_dispute'):
            QMessageBox.critical(update_form ,'Error', 'Nature of Dispute cannot be empty')
            return

        if not data.get('complainant'):
            QMessageBox.critical(update_form ,'Error', 'Complainant cannot be empty')
            return
        
        if not data.get('respondent'):
            QMessageBox.critical(update_form ,'Error', 'Respondent cannot be empty')
            return
        
        if not data.get('full_report'):
            QMessageBox.critical(update_form ,'Error', 'Report cannot be empty')
            return

        blotter.record_date = data['record_date']
        blotter.status = data['status']
        blotter.action_taken = data['action_taken']
        blotter.nature_of_dispute = data['nature_of_dispute']
        blotter.complainant = data['complainant']
        blotter.respondent = data['respondent']
        blotter.full_report = data['full_report']
        
        try:
            self.session.commit()
            QMessageBox.information(self.view, "Success", "Blotter updated successfully!")
        except Exception as e:
            self.session.rollback()  
            QMessageBox.critical(self.view, "Error", f"Failed to update blotter: {str(e)}")
        
        update_form.accept()   
        
    def browse(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Blotter to Browse')
            return
        if row_id:
            blotter = self.session.query(Blotter).get(row_id)
            
            if blotter:
                browse_form = BrowseBlotterForm()
                
                browse_form.set_fields(
                    record_date = blotter.record_date,
                    status = blotter.status,
                    action_taken = blotter.action_taken,
                    nature_of_dispute = blotter.nature_of_dispute,
                    complainant = blotter.complainant,
                    respondent = blotter.respondent,
                    full_report = blotter.full_report
                )
                
                browse_form.exec()
        
    
    def delete(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Blotter to Delete')
            return
        if row_id:
            blotter = self.session.query(Blotter).get(row_id)
            if blotter:
                reply = QMessageBox.question(
                    self.view,  # Parent widget
                    "Confirm Deletion",
                    f"Are you sure you want to delete '{str(blotter.full_text)[:20]}'?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
    
                if reply == QMessageBox.Yes:
                    self.session.delete(blotter)
                    self.session.commit()
                    self.refresh()
                    
class CertificateWindowController:
    def __init__(self, view: BaseWindow):
        self.db = Database()
        self.session = self.db.get_session()
        self.view = view
        self.refresh()
        
        self.view.btRefresh.clicked.connect(self.refresh)
        self.view.btSearch.clicked.connect(self.search)
        self.view.tbSearchBar.returnPressed.connect(self.search)
        self.view.btAdd.clicked.connect(self.add)
        self.view.btEdit.clicked.connect(self.edit)
        self.view.btDelete.clicked.connect(self.delete)
        self.view.btBrowse.clicked.connect(self.browse)
    
    def refresh(self):
        self.view.set_search_text('')
        certificates = self.session.query(Certificate).join(Certificate.resident).order_by(desc(Certificate.date_issued)).all()
        data = []
        for certificate in certificates:
            resident = certificate.resident 
            date_issued = certificate.date_issued
            type = certificate.type
            purpose = certificate.purpose
            resident_name = [
                getattr(resident, "first_name", ""),
                getattr(resident, "middle_name", ""),
                getattr(resident, "last_name", ""),
                getattr(resident, "suffix", "")
            ] if resident else []
    
            full_name = " ".join(filter(None, resident_name))
    
            result = [
                certificate.id,
                date_issued,
                type,
                full_name,
                purpose
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Date Issued', 'Type', 'Resident', 'Purpose'], data)
        
    def search(self):
        search_text = self.view.get_search_text().lower()
        search_terms = search_text.split()
    
        # Start with base query and join household
        query = self.session.query(Certificate).join(Certificate.resident)
    
        # Apply search filters
        if search_terms:
            conditions = []
            for term in search_terms:
                term_filter = or_(
                    Resident.first_name.ilike(f"%{term}%"),
                    Resident.middle_name.ilike(f"%{term}%"),
                    Resident.last_name.ilike(f"%{term}%"),
                    Resident.suffix.ilike(f"%{term}%"),
                    Certificate.date_issued.ilike(f"%{term}%"),
                    Certificate.type.ilike(f"%{term}%"),
                    Certificate.purpose.ilike(f"%{term}%")
                )
                conditions.append(term_filter)
            
            query = query.filter(and_(*conditions))
    
        # Apply sorting
        certificates = query.order_by(desc(Certificate.date_issued), Resident.last_name).all()
    
        # Format results
        data = []
        for certificate in certificates:
            resident = certificate.resident 
            date_issued = certificate.date_issued
            type = certificate.type
            purpose = certificate.purpose
            resident_name = [
                getattr(resident, "first_name", ""),
                getattr(resident, "middle_name", ""),
                getattr(resident, "last_name", ""),
                getattr(resident, "suffix", "")
            ] if resident else []
    
            full_name = " ".join(filter(None, resident_name))
    
            result = [
                certificate.id,
                date_issued,
                type,
                full_name,
                purpose
            ]
            data.append(result)
        
        self.view.load_table(['id', 'Date Issued', 'Type', 'Resident', 'Purpose'], data)
        
    def add(self):
        add_form = AddCertificateForm()
        
        add_form.form.tbDateIssued.setDate(QDate.currentDate())
        # Populate households for dropdown (household names and IDs)
        residents = self.session.query(Resident).order_by(Resident.first_name).all()
        resident_dict = {
            " ".join(filter(None, [
                resident.first_name,
                resident.middle_name,
                resident.last_name,
                resident.suffix
            ])): resident.id
            for resident in residents
        }
        
        # Set the fields for the form (populate the household dropdown with names)
        add_form.set_fields(name=list(resident_dict.keys()))
        add_form.form.cbName.setCurrentText('')  # Ensure no default value is set
    
        # Button signals
        add_form.addbar.btAdd.clicked.connect(lambda: self.on_add_button_click(add_form, resident_dict))
        add_form.addbar.btCancel.clicked.connect(add_form.reject)
    
        # Execute the form as a modal dialog
        add_form.exec()
    
    def on_add_button_click(self, add_form, resident_dict):
        data = add_form.get_fields()

        if not data.get('purpose'):
            QMessageBox.critical(add_form, 'Error', 'Purpose cannot be empty')
            return
    
        resident_name = data.get("name")
        if not resident_name:  
            QMessageBox.critical(add_form, "Error", "Please select a Resident.")
            return  
        
        resident_id = resident_dict.get(resident_name)
        if not resident_id:
            QMessageBox.critical(add_form, "Error", f"No matching Resident found for '{resident_name}'.")
            return  
            
        resident = self.session.query(Resident).get(resident_id)
    
        # Remove the 'household' key from the data dictionary since it's already handled
        del data['name']
    
        # Uppercase all string values in data (excluding non-string types)
        data = {key: value.upper() if isinstance(value, str) else value for key, value in data.items()}
    
        # Create a new Resident instance with the provided data
        new_certificate = Certificate(**data)
        new_certificate.resident = resident  # Associate the resident with the selected household
    
        try:
            # Add the new resident to the session and commit the transaction
            self.session.add(new_certificate)
            self.session.commit()
            QMessageBox.information(add_form, "Success", "Certificate added successfully!")
            self.refresh()  
            
            # Close the dialog only if everything is successful
            add_form.accept()
        except Exception as e:
            self.session.rollback()
            QMessageBox.critical(add_form, "Error", f"Failed to add resident: {str(e)}")
            add_form.activateWindow() 
            add_form.raise_()    
    
    def edit(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Certificate to Edit')
            return
        if row_id:
            certificate = self.session.query(Certificate).get(row_id)
    
            if certificate:
                update_form = UpdateCertificateForm()
    
                # Fetch households
                residents = self.session.query(Resident).order_by(Resident.first_name).all()
                resident_dict = {
                    " ".join(filter(None, [
                        resident.first_name,
                        resident.middle_name,
                        resident.last_name,
                        resident.suffix
                    ])): resident.id
                    for resident in residents
                }
                resident_names = list(resident_dict.keys())
    
                # Populate household combo box
                update_form.form.cbName.clear()
                update_form.form.cbName.addItems(resident_names)
                
                full_name = ""
                
                if Certificate.resident:
                    full_name = " ".join(filter(None, [
                        certificate.resident.first_name,
                        certificate.resident.middle_name,
                        certificate.resident.last_name,
                        certificate.resident.suffix
                    ]))
                    update_form.form.cbName.setCurrentText(full_name)
    
    
                # Set initial values
                update_form.set_fields(
                    name = full_name,
                    date_issued = certificate.date_issued,
                    type = certificate.type,
                    purpose = certificate.purpose
                )
    
                # Revert button restores original values
                update_form.updatebar.btRevert.clicked.connect(
                    lambda: update_form.set_fields(
                        name = full_name,
                        date_issued = certificate.date_issued,
                        type = certificate.type,
                        purpose = certificate.purpose
                    )
                )
    
                # Save button
                update_form.updatebar.btUpdate.clicked.connect(
                    lambda: self.save_update(certificate, update_form, resident_dict)
                )
    
                # Cancel button
                update_form.updatebar.btCancel.clicked.connect(update_form.reject)
    
                update_form.exec()
    
    
    def save_update(self, certificate, update_form, resident_dict):
        updated_data = update_form.get_fields()
    
        # Map household name to ID
        resident_name = update_form.form.cbName.currentText()
        resident_id = resident_dict.get(resident_name)

        if not updated_data.get('purpose'):
            QMessageBox.critical(update_form, 'Error', 'Purpose cannot be empty')
            return
    
        if not resident_name:
            QMessageBox.critical(update_form, "Error", "Please select a valid resident.")
            return
    
        resident = self.session.query(Resident).get(resident_id)
        if not resident:
            QMessageBox.critical(update_form, "Error", "Selected resident does not exist.")
            return
    
        # Update resident fields
        certificate.date_issued = updated_data['date_issued']
        certificate.type = updated_data['type']
        certificate.purpose = updated_data['purpose'].upper()
        certificate.resident = resident
    
        try:
            self.session.commit()
            QMessageBox.information(update_form, "Success", "Certificate updated successfully!")
            update_form.accept()
            self.refresh()
        except Exception as e:
            self.session.rollback()
            QMessageBox.critical(update_form, "Error", f"Failed to update Certificate: {str(e)}")
        
    def browse(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Certificate to Browse')
            return
        if row_id:
            certificate = self.session.query(Certificate).get(row_id)
            
            if certificate:
                browse_form = BrowseCertificateForm()
                browse_form.tbPrint.clicked.connect(lambda: self.export_certificate_to_docx(certificate, browse_form))
                resident = certificate.resident
                
                full_name = " ".join(filter(None, [
                    resident.first_name,
                    resident.middle_name,
                    resident.last_name,
                    resident.suffix
                ]))
                browse_form.set_fields(
                    name = full_name,
                    date_issued = certificate.date_issued,
                    type = certificate.type,
                    purpose = certificate.purpose
                   
                )
                
                browse_form.exec()
    
    def export_certificate_to_docx(self, certificate, form):
        resident = certificate.resident
        full_name = " ".join(filter(None, [
            resident.first_name,
            resident.middle_name,
            resident.last_name,
            resident.suffix
        ]))

        save_path, _ = QFileDialog.getSaveFileName(
            form,
            "Save Certificate",
            f"{certificate.date_issued}_{full_name}_{certificate.type}.docx",
            "Word Documents (*.docx)"
        )
        if not save_path:
            return

        try:
            if certificate.type == 'CLEARANCE':
                template_path = "certificates/Barangay-Clearance-Template.docx"
            elif certificate.type == 'INDIGENCY':
                template_path = "certificates/Indigency-Template.docx"
            elif certificate.type == 'RESIDENCY':
                template_path = "certificates/Residency-Template.docx"
            doc = Document(template_path)
        except Exception as e:
            QMessageBox.critical(form, "Error", f"Could not load DOCX template: {str(e)}")
            return

        # Resident info
        resident = certificate.resident
        name = " ".join(filter(None, [
            resident.first_name,
            resident.middle_name,
            resident.last_name,
            resident.suffix
        ]))
        age = self.calculate_age(resident.date_of_birth)
        date_str = certificate.date_issued.strftime("%dth day of %B, %Y")
        purpose = certificate.purpose
        cert_type = certificate.type

        # Replace placeholders
        replacements = {
            "{NAME}": name,
            "{AGE}": str(age),
            "{CIVIL}": resident.civil_status,
            "{CITIZEN}": resident.citizenship,
            "{DATE}": date_str,
            "{PURPOSE}": purpose
        }

        self.replace_placeholders_in_doc(doc, replacements)

        try:
            doc.save(save_path)
            QMessageBox.information(form, "Success", "Certificate saved successfully.")
        except Exception as e:
            QMessageBox.critical(form, "Error", f"Could not save DOCX file: {str(e)}")
    
    def calculate_age(self, date_of_birth: date) -> int:
        today = date.today()
        age = today.year - date_of_birth.year - (
            (today.month, today.day) < (date_of_birth.month, date_of_birth.day)
        )
        return age
    
    def replace_placeholders_in_doc(self, doc, replacements: dict):
        def process_runs(runs):
            # Combine all run texts into a single string
            full_text = "".join(run.text for run in runs)
            
            # Replace placeholders in the combined text
            for key, value in replacements.items():
                if key in full_text:
                    full_text = full_text.replace(key, value)
                    
                    # Update the first run with the new text and clear the rest
                    runs[0].text = full_text
                    for run in runs[1:]:
                        run.text = ""

        def process_paragraphs(paragraphs):
            for paragraph in paragraphs:
                process_runs(paragraph.runs)

        def process_table(table):
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)
                    for nested_table in cell.tables:  # Handle nested tables
                        process_table(nested_table)

        def process_section(section):
            # Process paragraphs in headers and footers
            for header in [section.header, section.footer]:
                process_paragraphs(header.paragraphs)
                for table in header.tables:
                    process_table(table)

        # Process paragraphs in the main document body
        process_paragraphs(doc.paragraphs)

        # Process tables in the main document body
        for table in doc.tables:
            process_table(table)

        # Process headers, footers, and other sections
        for section in doc.sections:
            process_section(section)
        
    def delete(self):
        try:
            row_id = self.view.get_table_row()
        except:
            QMessageBox.warning(self.view, 'Select Row', 'Please Select Certificate to Delete')
            return
        if row_id:
            certificate = self.session.query(Certificate).get(row_id)
            if certificate:
                reply = QMessageBox.question(
                    self.view, 
                    "Confirm Deletion",
                    f"Are you sure you want to delete Certificate: {certificate.date_issued} {certificate.type}: {certificate.resident.first_name} {certificate.resident.middle_name} {certificate.resident.last_name} {certificate.resident.suffix}?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
    
                if reply == QMessageBox.Yes:
                    self.session.delete(certificate)
                    self.session.commit()
                    self.refresh()

class AboutUsWindowController:
    def __init__(self, view: AboutWindow):
        self.db = Database()
        self.session = self.db.get_session()
        self.view = view
        self.load_data()

    def load_data(self):
        barangay = self.session.query(Barangay).first()
        users = self.session.query(User).join(User.resident).all()

        user_list = []

        if not users and not barangay:
            return

        for user in users:
            position = user.position
            resident = user.resident
            resident_name = [
                getattr(resident, "first_name", ""),
                getattr(resident, "middle_name", ""),
                getattr(resident, "last_name", ""),
                getattr(resident, "suffix", "")
            ] if resident else []
    
            full_name = " ".join(filter(None, resident_name))

            user_list.append({'name': full_name, 'position': position})
        
        custom_order = ["CAPTAIN", "SECRETARY", "TREASURER", "KAGAWAD", "TANOD"]

        self.view.load_data(
            name = barangay.name,
            history = barangay.history,
            mission = barangay.mission,
            vision = barangay.vision,
            members = sorted(user_list, key=lambda user: custom_order.index(user['position'].upper()) if user['position'].upper() in custom_order else len(custom_order))
        )
            
class SettingsWindowController:
    def __init__(self, view: SettingsWindow):
        self.view = view
        self.db = Database()
        self.session = self.db.get_session()
        self.load_data()

        self.view.btSave.clicked.connect(self.save_changes)
        self.view.btRevert.clicked.connect(self.revert_changes)
        self.view.btExport.clicked.connect(self.export_csv)
        self.view.btCreateBackup.clicked.connect(self.backup_database)
        self.view.btViewBackup.clicked.connect(self.switch_database)
    
    def load_data(self):
        barangay = self.session.query(Barangay).first()

        if barangay:
            self.view.set_fields(
                name = barangay.name,
                history = barangay.history,
                mission = barangay.mission,
                vision = barangay.vision
            )
    
    def save_changes(self):
        data = self.view.get_fields()
        data = {key: value.upper() if isinstance(value, str) else value for key, value in data.items()}

        barangay = self.session.query(Barangay).first()

        if not barangay:
            barangay = Barangay()
            barangay.name = data.get('name', '')
            barangay.history = data.get('history', '')
            barangay.mission = data.get('mission', '')
            barangay.vision = data.get('vision')

            try:
                # Add the new resident to the session and commit the transaction
                self.session.add(barangay)
                self.session.commit()
                QMessageBox.information(self.view, "Success", "Barangay added successfully!")
                self.load_data()  
                
            except Exception as e:
                self.session.rollback()
                QMessageBox.critical(self.view, "Error", f"Failed to add resident: {str(e)}")
        else:
            barangay.name = data.get('name', '')
            barangay.history = data.get('history', '')
            barangay.mission = data.get('mission', '')
            barangay.vision = data.get('vision')

            try:
                self.session.commit()
                QMessageBox.information(self.view, "Success", "Barangay updated successfully!")
                self.load_data()  
                
            except Exception as e:
                self.session.rollback()
                QMessageBox.critical(self.view, "Error", f"Failed to add resident: {str(e)}")


    def revert_changes(self):
        barangay = self.session.query(Barangay).first()

        if barangay:
            self.view.set_fields(
                name = barangay.name,
                history = barangay.history,
                mission = barangay.mission,
                vision = barangay.vision
            )

    def export_csv(self):
        timestamp = datetime.now().strftime("%Y_%m_%d")

        try:
            # Initialize a Pandas Excel writer to export multiple sheets
            file_path, _ = QFileDialog.getSaveFileName(
                self.view,
                "Save Excel File",
                f"export_{timestamp}.xlsx",
                "Excel Files (*.xlsx);;All Files (*)"
            )

            if not file_path:
                return  # User cancelled

            # Export Resident & Household data
            stmt_resident_household = (
                select(
                    Resident.id.label("resident_id"),
                    Resident.date_added,
                    Resident.first_name,
                    Resident.last_name,
                    Resident.middle_name,
                    Resident.suffix,
                    Resident.date_of_birth,
                    Resident.occupation,
                    Resident.civil_status,
                    Resident.citizenship,
                    Resident.sex,
                    Resident.education,
                    Resident.remarks,
                    Resident.phone1,
                    Resident.phone2,
                    Resident.email,
                    Resident.role,
                    Household.household_name,
                    Household.house_no,
                    Household.street,
                    Household.sitio,
                    Household.landmark
                )
                .join(Resident.household)
            )
            df_resident_household = pd.read_sql(stmt_resident_household, self.session.bind)

            # Export Resident & Certificate data
            stmt_resident_certificate = (
                select(
                    Resident.id.label("resident_id"),
                    Resident.first_name,
                    Resident.middle_name,
                    Resident.last_name,
                    Resident.suffix,
                    Certificate.type.label("certificate_type"),
                    Certificate.purpose.label("certificate_purpose"),
                    Certificate.date_issued
                )
                .join(Resident.certificates)
            )
            df_resident_certificate = pd.read_sql(stmt_resident_certificate, self.session.bind)

            # Export Blotter data
            stmt_blotter = select(Blotter)
            df_blotter = pd.read_sql(stmt_blotter, self.session.bind)

            # Export Resident & User data
            stmt_resident_user = (
                select(
                    Resident.id.label("resident_id"),
                    Resident.first_name,
                    Resident.last_name,
                    Resident.middle_name,
                    Resident.suffix,
                    Resident.date_of_birth,
                    Resident.occupation,
                    Resident.civil_status,
                    Resident.citizenship,
                    Resident.sex,
                    Resident.education,
                    Resident.remarks,
                    Resident.phone1,
                    Resident.phone2,
                    Resident.email,
                    Resident.role,
                    Household.household_name,
                    Household.house_no,
                    Household.street,
                    Household.sitio,
                    Household.landmark,
                    User.username,
                    User.position
                )
                .join(Resident.user)  
                .join(Resident.household)  
            )
            df_resident_user = pd.read_sql(stmt_resident_user, self.session.bind)

            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df_resident_household.to_excel(writer, sheet_name="RBI", index=False)
                df_resident_certificate.to_excel(writer, sheet_name="Certificates", index=False)
                df_blotter.to_excel(writer, sheet_name="Blotters", index=False)
                df_resident_user.to_excel(writer, sheet_name="Officials", index=False)

            QMessageBox.information(self.view, "Export Successful", f"Data saved to:\n{file_path}")

        except Exception as e:
            # Handle any errors
            QMessageBox.critical(self.view, "Export Failed", str(e))

    def backup_database(self):
        try:
            # Ask the user to select a folder to save the backup
            folder_path = QFileDialog.getExistingDirectory(
                self.view,
                "Select Folder to Save Backup",
                os.path.expanduser("~")  # Default to user's home directory
            )

            if not folder_path:
                return  # User cancelled

            # Get the current timestamp to append to the backup file name
            timestamp = datetime.now().strftime("%Y_%m_%d_%H_%M_%S")
            backup_filename = f"database_backup_{timestamp}.sqlite"

            # Specify the source (current database file) and the destination (selected folder)
            source_db_path = "main.sqlite"  # Adjust with the actual path
            destination_db_path = os.path.join(folder_path, backup_filename)

            # Copy the database to the selected folder
            shutil.copy(source_db_path, destination_db_path)

            # Show a message confirming the backup was successful
            QMessageBox.information(self.view, "Backup Successful", f"Database backed up to:\n{destination_db_path}")

        except Exception as e:
            # Handle any errors
            QMessageBox.critical(self.view, "Backup Failed", str(e))

    def switch_database(self):
        try:
            # Show confirmation dialog with extreme caution warning
            reply = QMessageBox.warning(
                self.view,
                "Extreme Caution!",
                "You are about to replace the current database with another one. "
                "This action cannot be undone! Are you sure you want to continue?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            
            # If the user selects 'No', exit the function
            if reply == QMessageBox.No:
                return

            # Ask the user to select the backup database (new database file)
            db_path, _ = QFileDialog.getOpenFileName(
                self.view,
                "Select SQLite Database to Replace Current DB",
                "",
                "SQLite Files (*.sqlite *.db *.SQLITE *.DB);;All Files (*)"

            )

            if not db_path:
                return  # User canceled the operation

            # Confirm with a second caution message before proceeding
            confirm = QMessageBox.warning(
                self.view,
                "Extreme Caution!",
                "This will overwrite the current database with the selected one. Are you sure?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )

            if confirm == QMessageBox.No:
                return  # If the user selects 'No', do nothing

            # Copy the selected new database to overwrite the current main.sqlite
            shutil.copy(db_path, 'main.sqlite')

            # Reinitialize the Database singleton to use the new database
            db_instance = Database(db_url='sqlite:///main.sqlite')  # Reuse the default database path
            db_instance.engine.dispose()  # Dispose of the old engine connection
            db_instance.engine = create_engine(f"sqlite:///{db_path}")  # Set the new database
            db_instance.Session = sessionmaker(bind=db_instance.engine)  # Rebind the session to the new engine

            # Notify the user
            QMessageBox.information(None, "Database Switched", "The database has been replaced successfully.")

        except Exception as e:
            # Handle any errors
            QMessageBox.critical(None, "Error", f"Failed to switch databases: {str(e)}")
    
class DashboardWindowController:
    def __init__(self, view: DashboardWindow):
        self.view = view
        self.db = Database()
        self.session = self.db.get_session()
        self.load_data()

    def update_bar_plot(self, plot, categories, values, title="Chart"):
        plot.clear()
        x = list(range(len(categories)))
        width = 0.6
        bars = pg.BarGraphItem(x=x, height=values, width=width, brush='skyblue')
        plot.addItem(bars)
        plot.setTitle(title)
        plot.getAxis('bottom').setTicks([list(zip(x, categories))])

        # Add value labels on top of bars
        for i, value in enumerate(values):
            label = pg.TextItem(html=f"<div style='text-align:center'>{value}</div>", anchor=(0.5, 1))
            label.setPos(x[i], value)
            plot.addItem(label)

    def load_data(self):
        # Count of entities
        households = self.session.query(Household).all()
        residents = self.session.query(Resident).all()
        blotters = self.session.query(Blotter).all()
        certificates = self.session.query(Certificate).all()

        self.update_bar_plot(
            self.view.plot_items[0],
            categories=['Households', 'Residents', 'Blotters', 'Certificates'],
            values=[len(households), len(residents), len(blotters), len(certificates)],
            title="Total Number of Entities Recorded"
        )

        # Sitio Count aggregate (Residents)
        stmt_sitio = select(Household.sitio).join(Resident.household)
        df = pd.read_sql(stmt_sitio, self.session.bind)
        sitio_counts = df["sitio"].value_counts().sort_index()
        categories = sitio_counts.index.tolist()
        values = sitio_counts.values.tolist()

        self.update_bar_plot(self.view.plot_items[1], categories, values, "Resident Sitio Distribution")

        # Age groups
        stmt_age = select(Resident.date_of_birth)
        df = pd.read_sql(stmt_age, self.session.bind)
        df["date_of_birth"] = pd.to_datetime(df["date_of_birth"])
        today = pd.Timestamp(datetime.today().date())
        df["age"] = (today - df["date_of_birth"]).dt.days // 365

        bins = [0, 18, 35, 50, 65, 120]
        labels = ["0-18", "19-35", "36-50", "51-65", "66+"]
        df["age_range"] = pd.cut(df["age"], bins=bins, labels=labels, right=True)

        age_counts = df["age_range"].value_counts().sort_index()
        categories = age_counts.index.tolist()
        values = age_counts.values.tolist()

        self.update_bar_plot(self.view.plot_items[2], categories, values, "Resident Age Group Distribution")

        # Sex
        stmt_sex = select(Resident.sex)
        df = pd.read_sql(stmt_sex, self.session.bind)

        sex_counts = df["sex"].value_counts().sort_index()
        categories = sex_counts.index.tolist()
        values = sex_counts.values.tolist()

        self.update_bar_plot(self.view.plot_items[3], categories, values, "Resident Gender/Sex Distribution")

        # Civil Status
        stmt_civil = select(Resident.civil_status)
        df = pd.read_sql(stmt_civil, self.session.bind)

        civil_counts = df["civil_status"].value_counts().sort_index()
        categories = civil_counts.index.tolist()
        values = civil_counts.values.tolist()

        self.update_bar_plot(self.view.plot_items[4], categories, values, "Resident Civil Status Distribution")

        # Blotter Status
        stmt_status = select(Blotter.status)
        df = pd.read_sql(stmt_status, self.session.bind)

        status_counts = df["status"].value_counts().sort_index()
        categories = status_counts.index.tolist()
        values = status_counts.values.tolist()

        self.update_bar_plot(self.view.plot_items[5], categories, values, "Blotter Status Distribution")


